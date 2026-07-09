"""
회의록 자동 생성 서비스 (회의록입력.ipynb 로직 통합)
LangChain + GPT를 사용하여 녹취록에서 회의록 자동 생성
"""
import os
import io
import re
import xml.etree.ElementTree as ET
from xml.dom import minidom
from typing import Dict, List, Tuple, Any, Optional
from pydantic import BaseModel, Field

from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import PydanticOutputParser

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# -------------------------------------------------------
# Pydantic 모델 정의
# -------------------------------------------------------
class MeetingAnalysis(BaseModel):
    """LLM이 반환할 회의록 데이터 구조"""
    agenda: str = Field(description="회의 안건 (명사형으로 간결하게)")
    summary: str = Field(description="전체 요약 및 회의 내용 (상세한 본문)")
    key_decisions: str = Field(description="결정 사항 및 회의 결과 (없으면 '없음')")
    action_items: str = Field(description="향후 계획, 할 일, 진행 일정")
    issues: str = Field(description="제기된 이슈, 우려 사항, 비고")
    next_agenda: str = Field(description="다음 회의 안건 (없으면 '없음')")
    keywords: str = Field(description="핵심 키워드 (쉼표 구분)")
    date: str = Field(description="회의 날짜 (YYYY-MM-DD)")


# -------------------------------------------------------
# 유틸리티 함수들
# -------------------------------------------------------
def clean_header(text: str) -> str:
    """헤더 정규화: 공백 제거 및 전각 기호 변환"""
    if not text:
        return ""
    text = "".join(text.split())
    text = text.replace("：", ":")
    return text


def format_text_content(text: str) -> str:
    """
    텍스트 포맷팅: 리스트 기호와 마침표 기준 줄바꿈 추가
    """
    if not text:
        return ""
    text = str(text).strip()

    formatted = text
    # 리스트 기호 기반 줄바꿈
    formatted = re.sub(r'(?<=\S)\s+-\s+', r'\n- ', formatted)
    formatted = re.sub(r'(?<=\S)\s+•\s+', r'\n• ', formatted)
    formatted = re.sub(r'(?<=\S)\s+(\d+\.)\s+', r'\n\1 ', formatted)

    # 긴 줄글 강제 분리 (마침표 기준)
    # 줄바꿈이 3개 미만이고 텍스트가 50자 이상이면 강제 분리
    if formatted.count('\n') < 3 and len(formatted) > 50:
        # 마침표 뒤에 공백이 있든 없든, 다음에 한글이나 영어 대문자가 오면 줄바꿈
        formatted = re.sub(r'(?<=[가-힣]{2})\.\s*(?=[가-힣A-Z])', r'.\n', formatted)

    return formatted


def delete_paragraph(paragraph):
    """문단을 문서에서 완전히 삭제"""
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def fill_element(element, text: str, label_prefix: Optional[str] = None):
    """Cell 또는 Paragraph에 텍스트 입력"""
    element.text = ""
    if not text:
        return

    is_cell = hasattr(element, 'paragraphs')

    if is_cell:
        p = element.paragraphs[0] if element.paragraphs else element.add_paragraph()
    else:
        p = element

    final_text = text
    if label_prefix:
        final_text = f"{label_prefix} : {text}"

    lines = final_text.split('\n')
    for i, line in enumerate(lines):
        clean_line = line.strip()
        if not clean_line:
            continue
        if i == 0:
            p.text = clean_line
        else:
            p.add_run('\n' + clean_line)

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if p.runs:
        for run in p.runs:
            run.font.size = Pt(10)


def extract_participants_from_text(stt_text: str) -> str:
    """녹취록 텍스트에서 화자 이름 추출"""
    if not stt_text:
        return ""
    speakers = set()
    pattern = re.compile(r'^\s*([^\s:].*?)\s*:\s+')

    for line in stt_text.split('\n'):
        match = pattern.match(line)
        if match:
            s = match.group(1).strip()
            if len(s) < 20:
                speakers.add(s)

    return ", ".join(sorted(list(speakers))) if speakers else ""


# -------------------------------------------------------
# 템플릿 자동 감지
# -------------------------------------------------------
def detect_template_type(template_path: str) -> int:
    """
    Word 템플릿 파일 분석하여 타입 감지
    Type 1: 기본형
    Type 2: 의견 중심형
    Type 3: 결과 중심형
    Type 4: 상세 관리형
    """
    if not os.path.exists(template_path):
        print(f"⚠️ 템플릿 파일 없음, 기본값(Type 1) 사용")
        return 1

    try:
        doc = Document(template_path)
        all_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text.replace(" ", "")
        for p in doc.paragraphs:
            all_text += p.text.replace(" ", "")

        if "의견사항" in all_text:
            return 2
        if "진행일정" in all_text or "이슈/비고" in all_text:
            return 4
        if "회의결과" in all_text:
            return 3
        return 1
    except Exception as e:
        print(f"⚠️ 템플릿 감지 실패: {e}")
        return 1


# -------------------------------------------------------
# LangChain 프롬프트 생성
# -------------------------------------------------------
def get_prompt_template(form_type: int) -> str:
    """양식 타입에 맞는 프롬프트 템플릿 반환"""
    base_instruction = """
당신은 전문 회의록 작성 비서입니다.
제공된 녹취록을 분석하여 요청된 JSON 형식에 맞춰 정보를 추출하세요.
내용이 없는 항목은 "없음"이라고 명시하세요.

**중요: 모든 항목은 반드시 개조식으로 작성하세요!**
- summary, issues, key_decisions, action_items, next_agenda는 각 항목을 "- "로 시작해야 합니다.
- 각 항목마다 반드시 줄바꿈(\n)을 넣어야 합니다.
- 긴 문장을 하나로 쓰지 말고, 항목별로 분리하세요.

예시:
summary: "- 첫 번째 논의 사항입니다.\n- 두 번째 논의 사항입니다.\n- 세 번째 결론입니다."
issues: "- 첫 번째 이슈입니다.\n- 두 번째 우려 사항입니다."
"""

    specific_instructions = {
        1: """
[중점 사항 - 한 줄 요약형 (Type 1)]
- 'summary' 항목에 회의의 전체 흐름, 주요 논의 사항, 결론을 작성하세요.
- 반드시 각 문장을 "- "로 시작하고 줄바꿈(\n)으로 구분하세요.
- 최소 5개 이상의 항목으로 작성하세요.
""",
        2: """
[중점 사항 - 의견 중심형 (Type 2)]
- 'summary': 각 사실을 "- "로 시작하여 작성
- 'issues': 각 의견을 "- "로 시작하여 작성
- 항목마다 반드시 줄바꿈(\n) 포함
""",
        3: """
[중점 사항 - 결과 중심형 (Type 3)]
- 'summary': 각 논의를 "- "로 시작하여 작성
- 'key_decisions': 각 결정을 "- "로 시작하여 작성
- 항목마다 반드시 줄바꿈(\n) 포함
""",
        4: """
[중점 사항 - 상세 관리형 (Type 4)]
- summary: 각 안건별 내용을 "- "로 시작하여 작성, 줄바꿈(\n)으로 구분
- issues: 각 이슈를 "- "로 시작하여 작성, 원인과 영향 포함
- key_decisions: 각 결정을 "- "로 시작하여 작성
- action_items: 각 항목을 "- [담당자] 할 일 (기한)" 형식으로 작성
- 모든 항목 사이에 반드시 줄바꿈(\n) 포함
"""
    }

    return f"{base_instruction}\n{specific_instructions.get(form_type, specific_instructions[1])}\n" + \
           "[형식 지침]\n{format_instructions}\n\n[녹취록]\n{transcript}"


# -------------------------------------------------------
# LangChain 회의록 분석
# -------------------------------------------------------
def analyze_transcript_with_langchain(
    api_key: str,
    transcript_text: str,
    form_type: int = 4
) -> Dict[str, Any]:
    """
    LangChain + GPT를 사용하여 녹취록 분석
    """
    try:
        llm = ChatOpenAI(
            model="gpt-4o-mini",
            temperature=0,
            openai_api_key=api_key
        )

        parser = PydanticOutputParser(pydantic_object=MeetingAnalysis)

        prompt = PromptTemplate(
            template=get_prompt_template(form_type),
            input_variables=["transcript"],
            partial_variables={"format_instructions": parser.get_format_instructions()}
        )

        chain = prompt | llm | parser

        # 텍스트 길이 제한 (토큰 제한 방지)
        result_obj = chain.invoke({"transcript": transcript_text[:15000]})

        # Pydantic v2 호환
        try:
            result_dict = result_obj.model_dump()
        except AttributeError:
            result_dict = result_obj.dict()

        return result_dict

    except Exception as e:
        print(f"❌ LangChain 분석 오류: {e}")
        return {
            "agenda": "분석 실패",
            "summary": "회의록 자동 생성 중 오류가 발생했습니다.",
            "key_decisions": "없음",
            "action_items": "없음",
            "issues": "없음",
            "next_agenda": "없음",
            "keywords": "",
            "date": ""
        }


# -------------------------------------------------------
# 서명부(참석자 명단) 처리
# -------------------------------------------------------
def process_signature_table(table, participants_str: str) -> bool:
    """표가 서명부인지 확인하고, 맞으면 참석자 명단 작성"""
    if not table.rows:
        return False
    try:
        if not table.rows[0].cells:
            return False
        headers = [clean_header(c.text) for c in table.rows[0].cells]
    except:
        return False

    is_sig = False
    req = ["서명", "Sign", "연락처"]
    if any(k in headers for k in req) and ("성명" in headers or "이름" in headers or "참석자" in headers):
        is_sig = True
    if not is_sig and "참석자" in headers and "성명" in headers and "소속" in headers:
        is_sig = True

    if is_sig:
        print("✍️ 서명부 발견! 명단 작성 중...")
        name_idx = -1
        for i, h in enumerate(headers):
            if "성명" in h or "이름" in h:
                name_idx = i
                break
        if name_idx == -1:
            for i, h in enumerate(headers):
                if "참석자" in h:
                    name_idx = i
                    break
        if name_idx == -1:
            return False

        people = [p.strip() for p in participants_str.split(',') if p.strip()]
        p_cursor = 0
        for r in range(1, len(table.rows)):
            row = table.rows[r]
            if len(row.cells) <= name_idx:
                continue
            if clean_header(row.cells[name_idx].text) in ["성명", "서명", "소속", "직급", "참석자", "비고"]:
                continue

            if p_cursor < len(people):
                fill_element(row.cells[name_idx], people[p_cursor])
                p_cursor += 1
            else:
                row.cells[name_idx].text = ""
        return True
    return False


# -------------------------------------------------------
# 문단 처리
# -------------------------------------------------------
def process_paragraphs(doc, data: Dict[str, str], header_map: Dict[str, str]):
    """일반 문단(본문) 스캔 및 수정"""
    print("🔎 일반 문단(본문) 스캔 및 수정 중...")
    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        text_raw = p.text.strip()
        if not text_raw:
            i += 1
            continue

        clean_txt = clean_header(text_raw)

        # 1. 혼합형 (제목 : 내용)
        if ":" in text_raw:
            parts = text_raw.split(":", 1)
            key_part = clean_header(parts[0])
            found_key = None
            for h_key, d_key in header_map.items():
                if h_key == key_part:
                    found_key = d_key
                    break

            if found_key:
                content = data.get(found_key, "")
                label = parts[0].strip()
                fill_element(p, content, label_prefix=label)
                print(f"  [문단] '{label}' 수정 완료")
                i += 1
                continue

        # 2. 제목형 (제목 다음 줄에 내용)
        found_key = None
        for h_key, d_key in header_map.items():
            if h_key == clean_txt:
                found_key = d_key
                break

        if found_key:
            content = format_text_content(data.get(found_key, ""))
            if i + 1 < len(doc.paragraphs):
                next_p = doc.paragraphs[i + 1]
                fill_element(next_p, content)
                print(f"  [문단] '{text_raw}' 하단 내용 작성 완료")

            # 잔여 텍스트 삭제
            check_idx = i + 2
            while check_idx < len(doc.paragraphs):
                check_p = doc.paragraphs[check_idx]
                check_txt = clean_header(check_p.text)
                is_next_header = False
                for hk in header_map.keys():
                    if hk == check_txt:
                        is_next_header = True
                        break
                if is_next_header:
                    break
                delete_paragraph(check_p)
            i += 1
        i += 1


# -------------------------------------------------------
# 메인 함수: Word 회의록 생성
# -------------------------------------------------------
def create_meeting_minutes_docx(
    transcript_data: List[Tuple[str, str]],
    speakers: List[str],
    file_info: Dict[str, Any],
    template_path: str,
    api_key: str,
    form_type: int = 4
) -> io.BytesIO:
    """
    LangGraph 워크플로우 기반 회의록 자동 생성
    복잡한 생성 과정을 구조화된 그래프로 관리

    Args:
        transcript_data: [(speaker, text), ...] 형식의 녹취록
        speakers: 화자 목록
        file_info: 파일 메타정보 (filename, created_at)
        template_path: Word 템플릿 파일 경로
        api_key: OpenAI API 키
        form_type: 템플릿 타입 (1~4)

    Returns:
        io.BytesIO: 생성된 Word 문서
    """
    # LangGraph 워크플로우 사용
    from .meeting_minutes_graph import generate_meeting_minutes_with_graph

    return generate_meeting_minutes_with_graph(
        transcript_data=transcript_data,
        speakers=speakers,
        file_info=file_info,
        template_path=template_path,
        api_key=api_key,
        form_type=form_type
    )

# 기존 코드는 meeting_minutes_graph.py로 이동됨

def _legacy_create_meeting_minutes_docx(
    transcript_data: List[Tuple[str, str]],
    speakers: List[str],
    file_info: Dict[str, Any],
    template_path: str,
    api_key: str,
    form_type: int = 4
) -> io.BytesIO:
    """
    레거시 버전 (LangGraph 없이)
    """
    print(f"🚀 회의록 자동 생성 시작...")
    print(f"✅ 선택된 양식 타입: Type {form_type}")

    # 2. 녹취록을 텍스트로 변환
    transcript_text = "\n".join([f"{spk}: {txt}" for spk, txt in transcript_data])

    # 3. LangChain 분석 실행
    print("🔗 LangChain 분석 중...")
    analysis_result = analyze_transcript_with_langchain(api_key, transcript_text, form_type)

    # 4. 메타 정보와 병합
    participants = ", ".join(speakers)
    data = {
        "DATE": analysis_result.get("date") or str(file_info.get("created_at", ""))[:10],
        "PARTICIPANTS": participants,
        "SUMMARY": analysis_result.get("summary", ""),
        "AGENDA": analysis_result.get("agenda", ""),
        "DECISIONS": analysis_result.get("key_decisions", ""),
        "ACTION_ITEMS": analysis_result.get("action_items", ""),
        "KEYWORDS": analysis_result.get("keywords", ""),
        "ISSUES": analysis_result.get("issues", ""),
        "NEXT_AGENDA": analysis_result.get("next_agenda", "")
    }

    # 5. Word 템플릿 로드
    print(f"📂 양식 파일 로드: {template_path}")
    doc = Document(template_path)

    # 헤더 매핑
    header_map = {
        "일시": "DATE", "날짜": "DATE", "회의일자": "DATE", "회의날짜": "DATE", "회의일시": "DATE",
        "참석자": "PARTICIPANTS", "회의참석자": "PARTICIPANTS",
        "회의안건": "AGENDA", "주제": "AGENDA", "회의주제": "AGENDA", "안건": "AGENDA",
        "내용": "SUMMARY", "회의내용": "SUMMARY", "주요안건및내용": "SUMMARY",
        "결과": "DECISIONS", "결정사항": "DECISIONS", "회의결과": "DECISIONS",
        "계획": "ACTION_ITEMS", "향후계획": "ACTION_ITEMS", "진행일정": "ACTION_ITEMS",
        "비고": "KEYWORDS", "이슈": "KEYWORDS", "의견사항": "ISSUES", "의견": "ISSUES",
        "다음회의": "NEXT_AGENDA"
    }

    column_map = {
        "회의내용": {"main": "SUMMARY", "이슈": "ISSUES", "비고": "ISSUES"},
        "결정사항": {"main": "DECISIONS", "진행일정": "ACTION_ITEMS", "계획": "ACTION_ITEMS"}
    }

    # 6. 표(Table) 처리
    for table in doc.tables:
        if process_signature_table(table, data["PARTICIPANTS"]):
            continue

        rows = table.rows
        if not rows:
            continue
        first_header = clean_header(rows[0].cells[0].text)

        # 멀티 컬럼 표 (양식 4)
        is_multi_column = False
        target_col_map = None
        for key, mapping in column_map.items():
            if key in first_header:
                header_cells_txt = [clean_header(c.text) for c in rows[0].cells]
                has_sub_col = False
                for txt in header_cells_txt[1:]:
                    if "내용" in txt or "Content" in txt:
                        has_sub_col = True
                    for sub_k in mapping.keys():
                        if sub_k != "main" and sub_k in txt:
                            has_sub_col = True
                if has_sub_col:
                    is_multi_column = True
                    target_col_map = mapping
                    break

        if is_multi_column and target_col_map:
            col_indices = {}
            for idx, cell in enumerate(rows[0].cells):
                txt = clean_header(cell.text)
                if "내용" in txt and "회의내용" not in txt:
                    col_indices[target_col_map["main"]] = idx
                for k, v in target_col_map.items():
                    if k != "main" and k in txt:
                        col_indices[v] = idx

            main_data = format_text_content(data.get(target_col_map["main"], "")).split('\n')
            main_data = [l.strip() for l in main_data if l.strip()]

            side_key = None
            side_data = []
            for k in col_indices:
                if k != target_col_map["main"]:
                    side_key = k
                    side_data = format_text_content(data.get(k, "")).split('\n')
                    side_data = [l.strip() for l in side_data if l.strip()]
                    break

            max_len = max(len(main_data), len(side_data))
            for i in range(max_len):
                target_row_idx = i + 1
                if target_row_idx >= len(rows):
                    break
                row = rows[target_row_idx]

                main_col_idx = col_indices.get(target_col_map["main"])
                if main_col_idx is not None:
                    if i < len(main_data):
                        fill_element(row.cells[main_col_idx], main_data[i])
                    else:
                        row.cells[main_col_idx].text = ""

                if side_key:
                    side_col_idx = col_indices.get(side_key)
                    if side_col_idx is not None:
                        if i < len(side_data):
                            fill_element(row.cells[side_col_idx], side_data[i])
                        else:
                            row.cells[side_col_idx].text = ""
            print(f"  - [표] '{first_header}' (멀티 컬럼) 작성 완료")
            continue

        # 일반 표 처리
        r_idx = 0
        while r_idx < len(rows):
            row = rows[r_idx]
            if not row.cells:
                r_idx += 1
                continue

            cell_text_raw = row.cells[0].text.strip()
            header_text = clean_header(cell_text_raw)

            is_mixed = False
            found_key = None
            label_prefix = None

            if ":" in cell_text_raw:
                parts = cell_text_raw.split(":", 1)
                key_part = clean_header(parts[0])
                if len(key_part) < 15:
                    for h_key, d_key in header_map.items():
                        if h_key in key_part:
                            found_key = d_key
                            is_mixed = True
                            label_prefix = parts[0].strip()
                            break

            if not is_mixed:
                if len(header_text) > 20 or header_text.startswith("-") or header_text.startswith("1."):
                    r_idx += 1
                    continue
                for h_key, d_key in header_map.items():
                    if h_key in header_text:
                        found_key = d_key
                        break

            if found_key:
                raw_content = data.get(found_key, "")
                formatted_content = format_text_content(raw_content)
                content_lines = [l.strip() for l in formatted_content.split('\n') if l.strip()]

                target_cell = None

                if is_mixed:
                    target_cell = row.cells[0]
                elif len(row.cells) > 1:
                    right_text = clean_header(row.cells[1].text)
                    if right_text not in ["담당자", "비고"] and "내용" not in right_text:
                        target_cell = row.cells[1]

                if target_cell is None:
                    if r_idx + 1 < len(rows):
                        bottom_text = clean_header(rows[r_idx + 1].cells[0].text)
                        is_next_header = False
                        if len(bottom_text) < 20:
                            for k in header_map.keys():
                                if k in bottom_text:
                                    is_next_header = True
                                    break
                        if not is_next_header:
                            target_cell = rows[r_idx + 1].cells[0]
                    else:
                        new_row = table.add_row()
                        target_cell = new_row.cells[0]

                if target_cell:
                    fill_element(target_cell, "\n".join(content_lines), label_prefix)
                    print(f"  - [표] '{header_text}' 작성 완료")
            r_idx += 1

    # 7. 문단 처리
    process_paragraphs(doc, data, header_map)

    # 8. BytesIO로 저장
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    print("✅ 회의록 생성 완료!")
    return output
