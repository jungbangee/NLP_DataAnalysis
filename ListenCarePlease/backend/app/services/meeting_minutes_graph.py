"""
LangGraph 기반 회의록 생성 워크플로우
복잡한 회의록 생성 과정을 구조화된 그래프로 관리
"""
import io
from typing import Dict, List, Tuple, Any, TypedDict
from langgraph.graph import StateGraph, END
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field

from docx import Document
from .meeting_minutes_service import (
    MeetingAnalysis,
    get_prompt_template,
    process_signature_table,
    process_paragraphs,
    fill_element,
    clean_header,
    format_text_content
)


# -------------------------------------------------------
# State 정의
# -------------------------------------------------------
class MeetingMinutesState(TypedDict):
    """회의록 생성 워크플로우의 상태"""
    # 입력 데이터
    transcript_data: List[Tuple[str, str]]
    speakers: List[str]
    file_info: Dict[str, Any]
    template_path: str
    api_key: str
    form_type: int

    # 중간 결과
    transcript_text: str
    analysis_result: Dict[str, Any]
    formatted_data: Dict[str, str]

    # 최종 결과
    output_docx: io.BytesIO
    error: str


# -------------------------------------------------------
# 노드 함수들
# -------------------------------------------------------
def prepare_transcript(state: MeetingMinutesState) -> MeetingMinutesState:
    """Step 1: 녹취록 텍스트 준비"""
    print("📝 [Step 1/4] 녹취록 텍스트 준비 중...")

    transcript_text = "\n".join([
        f"{speaker}: {text}"
        for speaker, text in state["transcript_data"]
    ])

    state["transcript_text"] = transcript_text
    print(f"  ✓ {len(state['transcript_data'])} 개 발화 준비 완료")
    return state


def analyze_with_llm(state: MeetingMinutesState) -> MeetingMinutesState:
    """Step 2: LangChain + GPT로 분석"""
    print("🤖 [Step 2/4] AI 분석 중...")

    try:
        llm = ChatOpenAI(
            model="gpt-4o-mini",
            temperature=0,
            openai_api_key=state["api_key"]
        )

        parser = PydanticOutputParser(pydantic_object=MeetingAnalysis)

        prompt = PromptTemplate(
            template=get_prompt_template(state["form_type"]),
            input_variables=["transcript"],
            partial_variables={"format_instructions": parser.get_format_instructions()}
        )

        chain = prompt | llm | parser

        # 텍스트 길이 제한 (토큰 제한 방지)
        result_obj = chain.invoke({"transcript": state["transcript_text"][:15000]})

        # Pydantic v2 호환
        try:
            result_dict = result_obj.model_dump()
        except AttributeError:
            result_dict = result_obj.dict()

        state["analysis_result"] = result_dict
        print(f"  ✓ AI 분석 완료 (안건: {result_dict.get('agenda', 'N/A')[:30]}...)")

    except Exception as e:
        print(f"  ✗ AI 분석 실패: {e}")
        state["error"] = f"AI analysis failed: {str(e)}"
        state["analysis_result"] = {
            "agenda": "분석 실패",
            "summary": "회의록 자동 생성 중 오류가 발생했습니다.",
            "key_decisions": "없음",
            "action_items": "없음",
            "issues": "없음",
            "next_agenda": "없음",
            "keywords": "",
            "date": ""
        }

    return state


def format_data(state: MeetingMinutesState) -> MeetingMinutesState:
    """Step 3: 데이터 포맷팅"""
    print("📋 [Step 3/4] 데이터 포맷팅 중...")

    participants = ", ".join(state["speakers"])
    analysis = state["analysis_result"]

    formatted_data = {
        "DATE": analysis.get("date") or str(state["file_info"].get("created_at", ""))[:10],
        "PARTICIPANTS": participants,
        "SUMMARY": analysis.get("summary", ""),
        "AGENDA": analysis.get("agenda", ""),
        "DECISIONS": analysis.get("key_decisions", ""),
        "ACTION_ITEMS": analysis.get("action_items", ""),
        "KEYWORDS": analysis.get("keywords", ""),
        "ISSUES": analysis.get("issues", ""),
        "NEXT_AGENDA": analysis.get("next_agenda", "")
    }

    state["formatted_data"] = formatted_data
    print("  ✓ 데이터 포맷팅 완료")
    return state


def generate_docx(state: MeetingMinutesState) -> MeetingMinutesState:
    """Step 4: Word 문서 생성"""
    print("📄 [Step 4/4] Word 문서 생성 중...")

    try:
        doc = Document(state["template_path"])
        data = state["formatted_data"]

        # 헤더 매핑
        header_map = {
            "일시": "DATE", "날짜": "DATE", "회의일자": "DATE", "회의날짜": "DATE", "회의일시": "DATE",
            "참석자": "PARTICIPANTS", "회의참석자": "PARTICIPANTS",
            "회의안건": "AGENDA", "주제": "AGENDA", "회의주제": "AGENDA", "안건": "AGENDA",
            "내용": "SUMMARY", "회의내용": "SUMMARY", "주요안건및내용": "SUMMARY",
            "결과": "DECISIONS", "결정사항": "DECISIONS", "회의결과": "DECISIONS",
            "계획": "ACTION_ITEMS", "향후계획": "ACTION_ITEMS", "진행일정": "ACTION_ITEMS",
            "비고": "KEYWORDS", "이슈": "KEYWORDS", "의견사항": "ISSUES", "의견": "ISSUES",
            "다음회의": "NEXT_AGENDA"
        }

        column_map = {
            "회의내용": {"main": "SUMMARY", "이슈": "ISSUES", "비고": "ISSUES"},
            "결정사항": {"main": "DECISIONS", "진행일정": "ACTION_ITEMS", "계획": "ACTION_ITEMS"}
        }

        # 표 처리
        for table in doc.tables:
            if process_signature_table(table, data["PARTICIPANTS"]):
                continue

            rows = table.rows
            if not rows:
                continue
            first_header = clean_header(rows[0].cells[0].text)

            # 멀티 컬럼 표 처리
            is_multi_column = False
            target_col_map = None
            for key, mapping in column_map.items():
                if key in first_header:
                    header_cells_txt = [clean_header(c.text) for c in rows[0].cells]
                    has_sub_col = False
                    for txt in header_cells_txt[1:]:
                        if "내용" in txt or "Content" in txt:
                            has_sub_col = True
                        for sub_k in mapping.keys():
                            if sub_k != "main" and sub_k in txt:
                                has_sub_col = True
                    if has_sub_col:
                        is_multi_column = True
                        target_col_map = mapping
                        break

            if is_multi_column and target_col_map:
                col_indices = {}
                for idx, cell in enumerate(rows[0].cells):
                    txt = clean_header(cell.text)
                    # 첫 번째 칼럼 (main) 감지: "회의내용", "결정사항" 등
                    if idx == 0 or (first_header in txt):
                        col_indices[target_col_map["main"]] = idx
                    # 두 번째 칼럼 (side) 감지: "이슈", "비고", "진행일정" 등
                    for k, v in target_col_map.items():
                        if k != "main" and k in txt:
                            col_indices[v] = idx

                main_data = format_text_content(data.get(target_col_map["main"], "")).split('\n')
                main_data = [l.strip() for l in main_data if l.strip()]

                side_key = None
                side_data = []
                for k in col_indices:
                    if k != target_col_map["main"]:
                        side_key = k
                        side_data = format_text_content(data.get(k, "")).split('\n')
                        side_data = [l.strip() for l in side_data if l.strip()]
                        break

                max_len = max(len(main_data), len(side_data))

                # 모든 데이터를 테이블에 넣기
                for i in range(max_len):
                    target_row_idx = i + 1
                    # 행이 부족하면 동적으로 추가
                    while target_row_idx >= len(table.rows):
                        table.add_row()
                    row = table.rows[target_row_idx]

                    main_col_idx = col_indices.get(target_col_map["main"])
                    if main_col_idx is not None:
                        if i < len(main_data):
                            fill_element(row.cells[main_col_idx], main_data[i])
                        else:
                            row.cells[main_col_idx].text = ""

                    if side_key:
                        side_col_idx = col_indices.get(side_key)
                        if side_col_idx is not None:
                            if i < len(side_data):
                                fill_element(row.cells[side_col_idx], side_data[i])
                            else:
                                row.cells[side_col_idx].text = ""

                # 빈 행 삭제 (모든 셀이 비어있는 행)
                rows_to_delete = []
                for row_idx in range(1, len(table.rows)):  # 헤더 행(0)은 제외
                    row = table.rows[row_idx]
                    # 모든 셀이 비어있는지 확인
                    is_empty = all(cell.text.strip() == "" for cell in row.cells)
                    if is_empty:
                        rows_to_delete.append(row)

                # 역순으로 삭제 (인덱스 오류 방지)
                for row in reversed(rows_to_delete):
                    tbl = table._element
                    tr = row._element
                    tbl.remove(tr)

                print(f"  [표] '{first_header}' (멀티 컬럼) 작성 완료 - {len(rows_to_delete)}개 빈 행 삭제")
                continue

            # 일반 표 처리 (간소화)
            r_idx = 0
            while r_idx < len(rows):
                row = rows[r_idx]
                if not row.cells:
                    r_idx += 1
                    continue

                cell_text_raw = row.cells[0].text.strip()
                header_text = clean_header(cell_text_raw)

                found_key = None
                for h_key, d_key in header_map.items():
                    if h_key in header_text:
                        found_key = d_key
                        break

                if found_key and len(row.cells) > 1:
                    raw_content = data.get(found_key, "")
                    formatted_content = format_text_content(raw_content)
                    fill_element(row.cells[1], formatted_content)

                r_idx += 1

        # 모든 표에서 빈 행 삭제 (헤더 행 제외)
        print("🧹 빈 행 정리 중...")
        for table in doc.tables:
            rows_to_delete = []
            for row_idx in range(1, len(table.rows)):  # 헤더 행(0)은 제외
                row = table.rows[row_idx]
                # 모든 셀이 비어있는지 확인
                is_empty = all(cell.text.strip() == "" for cell in row.cells)
                if is_empty:
                    rows_to_delete.append(row)

            # 역순으로 삭제
            for row in reversed(rows_to_delete):
                tbl = table._element
                tr = row._element
                tbl.remove(tr)

            if len(rows_to_delete) > 0:
                print(f"  ✓ 표에서 {len(rows_to_delete)}개 빈 행 삭제")

        # 문단 처리 (표 기반 템플릿에서는 불필요하므로 비활성화)
        # process_paragraphs(doc, data, header_map)

        # BytesIO로 저장
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        state["output_docx"] = output
        print("  ✓ Word 문서 생성 완료")

    except Exception as e:
        print(f"  ✗ Word 생성 실패: {e}")
        state["error"] = f"DOCX generation failed: {str(e)}"

    return state


# -------------------------------------------------------
# 그래프 구성
# -------------------------------------------------------
def create_meeting_minutes_workflow() -> StateGraph:
    """회의록 생성 워크플로우 그래프 생성"""
    workflow = StateGraph(MeetingMinutesState)

    # 노드 추가
    workflow.add_node("prepare", prepare_transcript)
    workflow.add_node("analyze", analyze_with_llm)
    workflow.add_node("format", format_data)
    workflow.add_node("generate", generate_docx)

    # 엣지 연결
    workflow.set_entry_point("prepare")
    workflow.add_edge("prepare", "analyze")
    workflow.add_edge("analyze", "format")
    workflow.add_edge("format", "generate")
    workflow.add_edge("generate", END)

    return workflow.compile()


# -------------------------------------------------------
# 메인 실행 함수
# -------------------------------------------------------
def generate_meeting_minutes_with_graph(
    transcript_data: List[Tuple[str, str]],
    speakers: List[str],
    file_info: Dict[str, Any],
    template_path: str,
    api_key: str,
    form_type: int = 4
) -> io.BytesIO:
    """
    LangGraph 워크플로우를 사용한 회의록 생성

    Args:
        transcript_data: [(speaker, text), ...] 형식의 녹취록
        speakers: 화자 목록
        file_info: 파일 메타정보
        template_path: Word 템플릿 경로
        api_key: OpenAI API 키
        form_type: 템플릿 타입 (1~4)

    Returns:
        io.BytesIO: 생성된 Word 문서
    """
    print("=" * 60)
    print("🚀 LangGraph 기반 회의록 생성 워크플로우 시작")
    print("=" * 60)

    # 초기 상태 설정
    initial_state: MeetingMinutesState = {
        "transcript_data": transcript_data,
        "speakers": speakers,
        "file_info": file_info,
        "template_path": template_path,
        "api_key": api_key,
        "form_type": form_type,
        "transcript_text": "",
        "analysis_result": {},
        "formatted_data": {},
        "output_docx": None,
        "error": ""
    }

    # 워크플로우 실행
    workflow = create_meeting_minutes_workflow()
    final_state = workflow.invoke(initial_state)

    print("=" * 60)
    print("✅ 워크플로우 완료")
    print("=" * 60)

    # 에러 체크
    if final_state.get("error"):
        raise Exception(f"Workflow error: {final_state['error']}")

    return final_state["output_docx"]
