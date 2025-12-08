import os
import io
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors
import urllib.request

# ----------------------------------------------------
# [유틸] 폰트 설정 (PDF용)
# ----------------------------------------------------
def register_korean_font():
    """PDF 생성을 위한 한글 폰트 등록 (나눔고딕)"""
    font_path = "NanumGothic.ttf"

    # 폰트 파일이 없으면 다운로드
    if not os.path.exists(font_path):
        try:
            print("📥 NanumGothic 폰트 다운로드 중...")
            url = "https://github.com/google/fonts/raw/main/ofl/nanumgothic/NanumGothic-Regular.ttf"
            urllib.request.urlretrieve(url, font_path)
        except Exception as e:
            print(f"⚠️ 폰트 다운로드 실패: {e}")
            return 'Helvetica'

    try:
        pdfmetrics.registerFont(TTFont('NanumGothic', font_path))
        return 'NanumGothic'
    except Exception as e:
        print(f"⚠️ 폰트 등록 실패: {e}")
        return 'Helvetica'

# ----------------------------------------------------
# [유틸] 워드 페이지 번호 함수
# ----------------------------------------------------
def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)

# ----------------------------------------------------
# 1. Word (DOCX) 생성
# ----------------------------------------------------
def create_docx(transcript_data, speakers, file_info):
    document = Document()

    # 기본 글꼴
    style = document.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # 제목
    heading = document.add_heading('회의 녹취록', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.size = Pt(20)
    heading.runs[0].font.bold = True
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # 기본 정보
    document.add_paragraph()
    info_table = document.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'

    file_name = file_info.get('filename', 'N/A')
    date_info = file_info.get('created_at', 'YYYY-MM-DD')
    speaker_list_str = ', '.join(speakers)

    info_data = [('회의 일시', str(date_info)), ('파일 이름', file_name), ('참여자', speaker_list_str)]

    for i, (header, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = header
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.1)

    # 대화 내용
    document.add_paragraph().add_run().add_break()
    document.add_heading('상세 대화 내용', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '화자'
    hdr[1].text = '내용'
    hdr[0].paragraphs[0].runs[0].bold = True
    hdr[0].width = Inches(0.8)
    hdr[1].width = Inches(5.8)

    for speaker, content in transcript_data:
        row = table.add_row().cells
        row[0].text = speaker
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = content
        row[0].width = Inches(0.8)
        row[1].width = Inches(5.8)

    # 페이지 번호
    add_page_number(document.sections[0].footer.paragraphs[0].add_run())
    
    # 메모리에 저장
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output

# ----------------------------------------------------
# 2. Excel (XLSX) 생성
# ----------------------------------------------------
def create_xlsx(transcript_data, file_info):
    # 1. 메타데이터 시트용 데이터
    meta_data = {
        "항목": ["회의 일시", "파일 이름", "생성일"],
        "내용": [
            str(file_info.get('created_at', 'N/A')),
            file_info.get('filename', 'N/A'),
            pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')
        ]
    }
    df_meta = pd.DataFrame(meta_data)

    # 2. 대화 내용 시트용 데이터
    df_content = pd.DataFrame(transcript_data, columns=["화자", "내용"])

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_meta.to_excel(writer, sheet_name='기본정보', index=False)
        df_content.to_excel(writer, sheet_name='대화내용', index=False)

        # 컬럼 너비 자동 조정
        worksheet = writer.sheets['대화내용']
        worksheet.column_dimensions['A'].width = 15
        worksheet.column_dimensions['B'].width = 80
    
    output.seek(0)
    return output

# ----------------------------------------------------
# 3. PDF 생성
# ----------------------------------------------------
def create_pdf(transcript_data, speakers, file_info):
    font_name = register_korean_font()
    
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4)
    elements = []

    # 스타일 정의
    styles = getSampleStyleSheet()
    style_title = ParagraphStyle('TitleKR', parent=styles['Title'], fontName=font_name, fontSize=20, leading=24)
    style_normal = ParagraphStyle('NormalKR', parent=styles['Normal'], fontName=font_name, fontSize=10, leading=14)
    style_header = ParagraphStyle('HeaderKR', parent=styles['Normal'], fontName=font_name, fontSize=11, leading=14, alignment=1)

    # 1. 제목
    elements.append(Paragraph("회의 녹취록", style_title))
    elements.append(Spacer(1, 20))

    # 2. 기본 정보 테이블
    file_name = file_info.get('filename', 'N/A')
    date_info = str(file_info.get('created_at', 'N/A'))
    speaker_str = ", ".join(speakers)

    info_data = [
        ["회의 일시", date_info],
        ["파일 이름", file_name],
        ["참여자", Paragraph(speaker_str, style_normal)]
    ]

    t_info = Table(info_data, colWidths=[100, 350])
    t_info.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), font_name),
        ('BACKGROUND', (0,0), (0,-1), colors.lightgrey),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-1), 1, colors.black),
        ('PADDING', (0,0), (-1,-1), 6),
    ]))
    elements.append(t_info)
    elements.append(Spacer(1, 30))

    # 3. 대화 내용 타이틀
    elements.append(Paragraph("상세 대화 내용", ParagraphStyle('SubTitle', parent=styles['Heading2'], fontName=font_name, alignment=1)))
    elements.append(Spacer(1, 10))

    # 4. 대화 내용 테이블
    table_data = [[Paragraph("화자", style_header), Paragraph("내용", style_header)]]

    for spk, txt in transcript_data:
        table_data.append([
            Paragraph(f"<b>{spk}</b>", style_normal),
            Paragraph(txt, style_normal)
        ])

    t_content = Table(table_data, colWidths=[80, 370], repeatRows=1)
    t_content.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), font_name),
        ('BACKGROUND', (0,0), (1,0), colors.whitesmoke),
        ('ALIGN', (0,0), (0,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('PADDING', (0,0), (-1,-1), 4),
    ]))
    elements.append(t_content)

    doc.build(elements)
    output.seek(0)
    return output


# ----------------------------------------------------
# 런타임 템플릿 생성
# ----------------------------------------------------
def create_runtime_template(output_path: str):
    """런타임에 기본 회의록 템플릿 생성"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 제목
    title = doc.add_heading('회의록', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 기본 정보 표
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'

    # 헤더 설정
    info_table.rows[0].cells[0].text = '회의 일시'
    info_table.rows[1].cells[0].text = '참석자'
    info_table.rows[2].cells[0].text = '회의 안건'

    # 회의 내용 표 (멀티 컬럼)
    doc.add_paragraph()
    doc.add_heading('회의 내용', level=2)

    content_table = doc.add_table(rows=6, cols=2)
    content_table.style = 'Table Grid'

    # 헤더
    content_table.rows[0].cells[0].text = '회의 내용'
    content_table.rows[0].cells[1].text = '이슈 / 비고'

    # 결정사항 표 (멀티 컬럼)
    doc.add_paragraph()
    doc.add_heading('결정 사항 및 향후 계획', level=2)

    decision_table = doc.add_table(rows=6, cols=2)
    decision_table.style = 'Table Grid'

    # 헤더
    decision_table.rows[0].cells[0].text = '결정 사항'
    decision_table.rows[0].cells[1].text = '진행 일정'

    # 저장
    doc.save(output_path)
    print(f"✅ 런타임 템플릿 생성 완료: {output_path}")
