#!/usr/bin/env python3
"""서버에서 호출: python generate_docx.py <json_file> <output_docx>"""
import sys, json, io, base64
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

data = json.load(open(sys.argv[1], encoding='utf-8'))
out  = sys.argv[2]
mode = data.get('mode', 'analysis')

doc = Document()

# 기본 스타일
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Malgun Gothic'
    return p

def add_para(text='', bold=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def set_cell_bg(cell, hex_color):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=10, color=None, align=None):
    """셀에 텍스트 설정 (기존 내용 초기화)"""
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_header_row(tbl, headers, bg='2D5F8A'):
    """헤더 행 추가 (배경색 + 흰색 볼드)"""
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=10, color=(255,255,255))
        set_cell_bg(hdr[i], bg)

# ── 레이더 차트 이미지를 흰 배경으로 변환 ────────────────────────
def radar_img_white_bg(img_b64):
    """canvas toDataURL은 다크 배경이므로 PIL로 흰 배경 합성"""
    try:
        from PIL import Image
        img_bytes = base64.b64decode(img_b64)
        img = Image.open(io.BytesIO(img_bytes)).convert('RGBA')
        bg = Image.new('RGBA', img.size, (255, 255, 255, 255))
        bg.paste(img, mask=img.split()[3])
        out_buf = io.BytesIO()
        bg.convert('RGB').save(out_buf, format='PNG')
        out_buf.seek(0)
        return out_buf
    except ImportError:
        # PIL 없으면 원본 그대로
        return io.BytesIO(base64.b64decode(img_b64))

# ════════════════════════════════════════════════════════════════
if mode == 'analysis':
    instructor = data.get('instructor','')
    date       = data.get('date','')
    course     = data.get('course_id','')

    add_heading(f'{instructor} 강의 분석 리포트', 1)
    add_para(f'📅 {date}   🏷 {course}', size=11)
    doc.add_paragraph()

    # 카테고리 요약
    cats_summary = data.get('cats_summary', [])
    if cats_summary:
        add_heading('카테고리별 점수', 2)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        add_header_row(tbl, ['카테고리', '점수', '달성률'])
        for c in cats_summary:
            row = tbl.add_row().cells
            set_cell_text(row[0], c.get('name',''))
            set_cell_text(row[1], f"{c.get('score',0):.1f} / 5.0", bold=True)
            set_cell_text(row[2], f"{c.get('pct',0):.0f}%", bold=True)
        doc.add_paragraph()

    # 차트 이미지
    for img_key, title in [('radar_img','카테고리별 달성률 (레이더)'), ('bar_img','세부 항목 점수')]:
        img_b64 = data.get(img_key)
        if img_b64:
            add_heading(title, 2)
            img_buf = radar_img_white_bg(img_b64)
            doc.add_picture(img_buf, width=Inches(5))
            doc.add_paragraph()

    # 카테고리 상세
    details = data.get('details', [])
    for cat in details:
        add_heading(cat.get('title',''), 2)
        for item in cat.get('items', []):
            p = doc.add_paragraph()
            r = p.add_run(f"{item.get('name','')}  {item.get('score','')}/5")
            r.font.name = 'Malgun Gothic'; r.bold = True
            if item.get('reason'):
                add_para(item['reason'], size=9)
            if item.get('example'):
                add_para(f"💡 {item['example']}", size=9)
            doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
elif mode == 'instructor':
    name     = data.get('instructor','')
    subtitle = data.get('subtitle','').strip()

    add_heading(f'{name} 강사 종합평가 리포트', 1)

    # ── 1. 날짜·강의수·등급 → 표로 출력 ────────────────────────
    # subtitle 예: "2026-02-02 ~ 2026-02-02 · 1회 강의  미흡"
    parts = [p.strip() for p in subtitle.split('  ') if p.strip()]
    info_rows = []
    for part in parts:
        if '~' in part and '강의' in part:
            # "2026-02-02 ~ 2026-02-02 · 1회 강의" 파싱
            if '·' in part:
                date_part, cnt_part = part.split('·', 1)
                info_rows.append(('기간', date_part.strip()))
                info_rows.append(('강의 횟수', cnt_part.strip()))
            else:
                info_rows.append(('기간', part))
        elif part:
            info_rows.append(('평가 등급', part))

    if info_rows:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        add_header_row(tbl, ['항목', '내용'])
        for label, value in info_rows:
            row = tbl.add_row().cells
            set_cell_text(row[0], label, bold=True)
            set_cell_text(row[1], value)
        doc.add_paragraph()
    else:
        add_para(subtitle, size=11)
        doc.add_paragraph()

    # ── 2. 종합 통계 (빈 헤더 행 제거) ─────────────────────────
    stats = data.get('stats_cards', [])
    if stats:
        add_heading('종합 통계', 2)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        add_header_row(tbl, ['항목', '값'])
        for s in stats:
            label = s.get('label','')
            value = str(s.get('value',''))
            if not label and not value:
                continue
            row = tbl.add_row().cells
            set_cell_text(row[0], label)
            set_cell_text(row[1], value, bold=True)
        doc.add_paragraph()

    # ── 3. 역량 레이더 (흰 배경 합성) ───────────────────────────
    img_b64 = data.get('radar_img')
    if img_b64:
        add_heading('역량 레이더', 2)
        img_buf = radar_img_white_bg(img_b64)
        doc.add_picture(img_buf, width=Inches(4))
        doc.add_paragraph()

    # ── 4. 카테고리별 평균 (빈 헤더 행 제거) ────────────────────
    cat_bars = data.get('cat_bars', [])
    if cat_bars:
        add_heading('카테고리별 평균', 2)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        add_header_row(tbl, ['카테고리', '평균 점수'])
        for b in cat_bars:
            name_b = b.get('name','')
            score  = str(b.get('score',''))
            if not name_b and not score:
                continue
            row = tbl.add_row().cells
            set_cell_text(row[0], name_b)
            set_cell_text(row[1], score, bold=True)
        doc.add_paragraph()

    # 텍스트 섹션
    for key, title in [('profile_summary','역량 프로파일'),('trajectory_note','추세 분석'),('consistency_note','일관성')]:
        txt = data.get(key,'')
        if txt:
            add_heading(title, 2)
            add_para(txt)
            doc.add_paragraph()

    # 강점/약점/목표
    for key, title in [('systematic_strengths','체계적 강점'),('systematic_weaknesses','체계적 약점'),('development_goals','개발 과제')]:
        items = data.get(key, [])
        if items:
            add_heading(title, 2)
            for item in items:
                add_para(f'• {item}')
            doc.add_paragraph()

    # 강의 이력
    lectures = data.get('lectures', [])
    if lectures:
        add_heading('강의 이력', 2)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        add_header_row(tbl, ['날짜', '과목', '카테고리 점수'])
        for lec in lectures:
            row = tbl.add_row().cells
            set_cell_text(row[0], lec.get('date',''))
            set_cell_text(row[1], lec.get('course_id',''))
            set_cell_text(row[2], lec.get('scores',''))
        doc.add_paragraph()

doc.save(out)
print('OK')